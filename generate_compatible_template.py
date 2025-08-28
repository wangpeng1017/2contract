#!/usr/bin/env python3
"""
生成系统兼容的Word模板文件
基于参考文档内容，创建标准占位符格式的模板
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_compatible_template():
    """创建系统兼容的Word模板"""
    
    print("📝 创建系统兼容的Word模板")
    print("=" * 50)
    
    # 创建新文档
    doc = Document()
    
    # 设置文档样式
    setup_document_styles(doc)
    
    # 添加标题
    add_title(doc)
    
    # 添加合同基本信息
    add_contract_info(doc)
    
    # 添加甲乙双方信息
    add_parties_info(doc)
    
    # 添加合同正文
    add_contract_body(doc)
    
    # 添加签署信息
    add_signature_section(doc)
    
    # 保存文档
    output_path = "系统兼容-采购合同模板.docx"
    doc.save(output_path)
    
    print(f"✅ 模板已生成: {output_path}")
    
    # 验证模板
    verify_template(output_path)
    
    return output_path

def setup_document_styles(doc):
    """设置文档样式"""
    
    # 设置正文样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 创建标题样式
    if 'Title Custom' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('Title Custom', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = '黑体'
        title_font.size = Pt(18)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_title(doc):
    """添加标题"""
    
    title = doc.add_paragraph()
    title.style = 'Title Custom'
    title.add_run('采购合同')
    
    # 添加空行
    doc.add_paragraph()

def add_contract_info(doc):
    """添加合同基本信息"""
    
    # 合同编号
    p1 = doc.add_paragraph()
    p1.add_run('合同编号：').bold = True
    p1.add_run('{{合同编号}}')
    
    # 合同类型
    p2 = doc.add_paragraph()
    p2.add_run('合同类型：').bold = True
    p2.add_run('{{合同类型}}')
    
    # 签署日期
    p3 = doc.add_paragraph()
    p3.add_run('签署日期：').bold = True
    p3.add_run('{{签署日期}}')
    
    # 添加空行
    doc.add_paragraph()

def add_parties_info(doc):
    """添加甲乙双方信息"""
    
    # 甲方信息
    doc.add_paragraph().add_run('甲方（采购方）信息：').bold = True
    
    p1 = doc.add_paragraph()
    p1.add_run('公司名称：').bold = True
    p1.add_run('{{甲方公司名称}}')
    
    p2 = doc.add_paragraph()
    p2.add_run('联系人：').bold = True
    p2.add_run('{{甲方联系人}}')
    
    p3 = doc.add_paragraph()
    p3.add_run('联系电话：').bold = True
    p3.add_run('{{甲方电话}}')
    
    p4 = doc.add_paragraph()
    p4.add_run('电子邮箱：').bold = True
    p4.add_run('{{联系邮箱}}')
    
    # 添加空行
    doc.add_paragraph()
    
    # 乙方信息
    doc.add_paragraph().add_run('乙方（供应方）信息：').bold = True
    
    p5 = doc.add_paragraph()
    p5.add_run('公司名称：').bold = True
    p5.add_run('{{乙方公司名称}}')
    
    p6 = doc.add_paragraph()
    p6.add_run('联系人：').bold = True
    p6.add_run('{{乙方联系人}}')
    
    p7 = doc.add_paragraph()
    p7.add_run('联系电话：').bold = True
    p7.add_run('{{乙方电话}}')
    
    # 添加空行
    doc.add_paragraph()

def add_contract_body(doc):
    """添加合同正文"""
    
    # 前言
    preamble = doc.add_paragraph()
    preamble.add_run('根据《中华人民共和国合同法》及相关法律法规，甲乙双方在平等、自愿、公平、诚信的基础上，就甲方向乙方采购货物事宜，经友好协商，达成如下协议：')
    
    doc.add_paragraph()
    
    # 第一条 货物信息
    doc.add_paragraph().add_run('第一条 货物信息').bold = True
    
    p1_1 = doc.add_paragraph()
    p1_1.add_run('1.1 货物名称：')
    p1_1.add_run('{{产品清单}}')
    
    p1_2 = doc.add_paragraph()
    p1_2.add_run('1.2 规格型号：')
    p1_2.add_run('{{产品规格}}')
    
    p1_3 = doc.add_paragraph()
    p1_3.add_run('1.3 数量：')
    p1_3.add_run('{{产品数量}}')
    
    p1_4 = doc.add_paragraph()
    p1_4.add_run('1.4 质量标准：')
    p1_4.add_run('{{质量标准}}')
    
    doc.add_paragraph()
    
    # 第二条 价格条款
    doc.add_paragraph().add_run('第二条 价格条款').bold = True
    
    p2_1 = doc.add_paragraph()
    p2_1.add_run('2.1 合同总金额：')
    p2_1.add_run('{{合同金额}}')
    
    p2_2 = doc.add_paragraph()
    p2_2.add_run('2.2 价格包含：货物价格、包装费、运输费等所有费用')
    
    doc.add_paragraph()
    
    # 第三条 交付条款
    doc.add_paragraph().add_run('第三条 交付条款').bold = True
    
    p3_1 = doc.add_paragraph()
    p3_1.add_run('3.1 交付时间：')
    p3_1.add_run('{{交付时间}}')
    
    p3_2 = doc.add_paragraph()
    p3_2.add_run('3.2 交付地点：')
    p3_2.add_run('{{交付地点}}')
    
    p3_3 = doc.add_paragraph()
    p3_3.add_run('3.3 验收标准：')
    p3_3.add_run('{{验收标准}}')
    
    doc.add_paragraph()
    
    # 第四条 付款方式
    doc.add_paragraph().add_run('第四条 付款方式').bold = True
    
    p4_1 = doc.add_paragraph()
    p4_1.add_run('4.1 付款方式：')
    p4_1.add_run('{{付款方式}}')
    
    p4_2 = doc.add_paragraph()
    p4_2.add_run('4.2 付款期限：')
    p4_2.add_run('{{付款期限}}')
    
    doc.add_paragraph()
    
    # 第五条 质量保证
    doc.add_paragraph().add_run('第五条 质量保证').bold = True
    
    p5_1 = doc.add_paragraph()
    p5_1.add_run('5.1 是否包含保险：')
    p5_1.add_run('{{是否包含保险}}')
    
    p5_2 = doc.add_paragraph()
    p5_2.add_run('5.2 质量问题处理：乙方应对货物质量负责，如发现质量问题，乙方应及时处理')
    
    doc.add_paragraph()
    
    # 第六条 违约责任
    doc.add_paragraph().add_run('第六条 违约责任').bold = True
    
    p6_1 = doc.add_paragraph()
    p6_1.add_run('6.1 甲方违约责任：甲方未按约定时间付款的，应承担违约责任')
    
    p6_2 = doc.add_paragraph()
    p6_2.add_run('6.2 乙方违约责任：乙方未按约定时间交付货物或货物质量不符合要求的，应承担违约责任')
    
    p6_3 = doc.add_paragraph()
    p6_3.add_run('6.3 违约金：违约方应向守约方支付合同总金额的5%作为违约金')
    
    doc.add_paragraph()
    
    # 第七条 争议解决
    doc.add_paragraph().add_run('第七条 争议解决').bold = True
    
    p7_1 = doc.add_paragraph()
    p7_1.add_run('7.1 本合同履行过程中发生的争议，双方应友好协商解决')
    
    p7_2 = doc.add_paragraph()
    p7_2.add_run('7.2 协商不成的，可向合同签署地人民法院起诉')
    
    doc.add_paragraph()
    
    # 第八条 其他约定
    doc.add_paragraph().add_run('第八条 其他约定').bold = True
    
    p8_1 = doc.add_paragraph()
    p8_1.add_run('8.1 特别约定：')
    p8_1.add_run('{{特别约定}}')
    
    p8_2 = doc.add_paragraph()
    p8_2.add_run('8.2 本合同一式两份，甲乙双方各执一份，具有同等法律效力')
    
    p8_3 = doc.add_paragraph()
    p8_3.add_run('8.3 本合同自双方签字盖章之日起生效')

def add_signature_section(doc):
    """添加签署部分"""
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 签署区域
    signature_table = doc.add_table(rows=3, cols=2)
    signature_table.style = 'Table Grid'
    
    # 第一行：甲方乙方
    signature_table.cell(0, 0).text = '甲方（盖章）：'
    signature_table.cell(0, 1).text = '乙方（盖章）：'
    
    # 第二行：法定代表人
    signature_table.cell(1, 0).text = '法定代表人：'
    signature_table.cell(1, 1).text = '法定代表人：'
    
    # 第三行：签署日期
    signature_table.cell(2, 0).text = '签署日期：{{签署日期}}'
    signature_table.cell(2, 1).text = '签署日期：{{签署日期}}'

def verify_template(template_path):
    """验证模板"""
    
    print(f"\n🔍 验证模板: {template_path}")
    
    if not os.path.exists(template_path):
        print("❌ 模板文件不存在")
        return False
    
    file_size = os.path.getsize(template_path)
    print(f"📄 文件大小: {file_size:,} 字节")
    
    # 检查占位符
    import zipfile
    import tempfile
    
    with tempfile.TemporaryDirectory() as temp_dir:
        with zipfile.ZipFile(template_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        with open(document_xml_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
        
        print(f"📄 XML长度: {len(xml_content):,} 字符")
        
        # 查找占位符
        import re
        placeholders = re.findall(r'\{\{([^}]+)\}\}', xml_content)
        unique_placeholders = list(set(placeholders))
        
        print(f"🎯 找到占位符: {len(unique_placeholders)} 个")
        for placeholder in sorted(unique_placeholders):
            print(f"  - {{{{ {placeholder} }}}}")
        
        # 检查系统支持的13个字段
        system_fields = [
            "甲方公司名称", "乙方公司名称", "合同类型", "合同金额", "签署日期",
            "甲方联系人", "甲方电话", "乙方联系人", "联系邮箱", "付款方式",
            "产品清单", "是否包含保险", "特别约定"
        ]
        
        supported_fields = [field for field in system_fields if field in unique_placeholders]
        print(f"\n✅ 系统支持字段: {len(supported_fields)}/13 个")
        for field in supported_fields:
            print(f"  ✓ {field}")
        
        missing_fields = [field for field in system_fields if field not in unique_placeholders]
        if missing_fields:
            print(f"\n⚠️  缺少字段: {len(missing_fields)} 个")
            for field in missing_fields:
                print(f"  - {field}")
    
    return True

def main():
    """主函数"""
    
    print("📝 系统兼容Word模板生成器")
    print("基于参考文档创建标准占位符格式的模板")
    print("=" * 60)
    
    try:
        template_path = create_compatible_template()
        
        print(f"\n🎉 模板生成成功！")
        print(f"📁 文件路径: {os.path.abspath(template_path)}")
        print(f"📋 特点:")
        print(f"  - 使用标准 {{{{字段名}}}} 占位符格式")
        print(f"  - 简化的段落结构，避免复杂表格")
        print(f"  - 包含完整的合同条款和业务逻辑")
        print(f"  - 优化的XML结构，确保系统兼容性")
        print(f"  - 支持系统预定义的13个字段")
        
    except Exception as e:
        print(f"❌ 生成模板时出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
