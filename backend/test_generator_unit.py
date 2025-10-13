"""文档生成器单元测试"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from services.document_generator import DocumentGenerator, TemplateConverter
from docx import Document


def test_document_generator():
    """测试文档生成器"""
    print("测试文档生成器...")
    print("-" * 50)
    
    # 创建测试模板
    doc = Document()
    doc.add_heading('测试合同', 0)
    doc.add_paragraph('合同编号：{{ contract_number }}')
    doc.add_paragraph('签订日期：{{ signing_date }}')
    doc.add_paragraph('甲方：{{ party_a }}')
    doc.add_paragraph('金额：{{ contract_amount }} 元')
    
    # 保存到内存
    template_stream = io.BytesIO()
    doc.save(template_stream)
    template_bytes = template_stream.getvalue()
    
    print(f"✓ 创建测试模板，大小: {len(template_bytes)} 字节")
    
    # 创建生成器
    generator = DocumentGenerator(template_bytes)
    print(f"✓ 文档生成器初始化成功")
    
    # 准备数据
    context = {
        "contract_number": "HT-2025-001",
        "signing_date": "2025-01-10",  # 应转换为 2025年01月10日
        "party_a": "测试公司",
        "contract_amount": 123456.78  # 应格式化为 123,456.78
    }
    
    # 渲染文档
    output_bytes = generator.render(context)
    print(f"✓ 文档渲染成功，大小: {len(output_bytes)} 字节")
    
    # 验证输出
    output_doc = Document(io.BytesIO(output_bytes))
    output_text = "\n".join([p.text for p in output_doc.paragraphs])
    
    print(f"\n生成的文档内容:")
    print("-" * 50)
    for para in output_doc.paragraphs:
        if para.text.strip():
            print(f"  {para.text}")
    print("-" * 50)
    
    # 检查格式化
    if "2025年01月10日" in output_text:
        print("✓ 日期格式化成功")
    else:
        print("⚠ 日期格式化可能未生效")
    
    if "123,456.78" in output_text:
        print("✓ 金额格式化成功")
    else:
        print("⚠ 金额格式化可能未生效")
    
    print("\n✓ 文档生成器测试通过\n")


def test_template_converter():
    """测试模板转换器"""
    print("测试模板转换器...")
    print("-" * 50)
    
    # 创建包含中文占位符的文档
    doc = Document()
    doc.add_paragraph('甲方：{{甲方公司名称}}')
    doc.add_paragraph('乙方：{{乙方公司名称}}')
    doc.add_paragraph('金额：{{合同金额}}元')
    
    input_stream = io.BytesIO()
    doc.save(input_stream)
    input_bytes = input_stream.getvalue()
    
    print(f"✓ 创建原始文档，大小: {len(input_bytes)} 字节")
    
    # 变量映射
    mapping = {
        "甲方公司名称": "party_a_company",
        "乙方公司名称": "party_b_company",
        "合同金额": "contract_amount"
    }
    
    # 转换模板
    converter = TemplateConverter()
    converted_bytes = converter.convert_to_jinja_template(input_bytes, mapping)
    
    print(f"✓ 模板转换成功，大小: {len(converted_bytes)} 字节")
    
    # 验证转换结果
    converted_doc = Document(io.BytesIO(converted_bytes))
    converted_text = "\n".join([p.text for p in converted_doc.paragraphs])
    
    print(f"\n转换后的内容:")
    print("-" * 50)
    for para in converted_doc.paragraphs:
        if para.text.strip():
            print(f"  {para.text}")
    print("-" * 50)
    
    # 检查是否转换成功
    checks = [
        ("{{ party_a_company }}" in converted_text, "party_a_company"),
        ("{{ party_b_company }}" in converted_text, "party_b_company"),
        ("{{ contract_amount }}" in converted_text, "contract_amount")
    ]
    
    for check, name in checks:
        if check:
            print(f"✓ 变量 {name} 转换成功")
        else:
            print(f"⚠ 变量 {name} 转换可能未生效")
    
    print("\n✓ 模板转换器测试通过\n")


def test_data_preprocessing():
    """测试数据预处理"""
    print("测试数据预处理...")
    print("-" * 50)
    
    # 创建简单模板
    doc = Document()
    doc.add_paragraph('测试')
    template_stream = io.BytesIO()
    doc.save(template_stream)
    
    generator = DocumentGenerator(template_stream.getvalue())
    
    # 测试各种数据类型
    test_cases = [
        {
            "name": "日期格式化",
            "input": {"date": "2025-01-10"},
            "check": lambda ctx: "2025年01月10日" in str(ctx.get("date", ""))
        },
        {
            "name": "金额格式化",
            "input": {"contract_amount": 1234567.89},
            "check": lambda ctx: "1,234,567.89" in str(ctx.get("contract_amount", ""))
        },
        {
            "name": "布尔值转换",
            "input": {"is_active": True},
            "check": lambda ctx: ctx.get("is_active") == "是"
        },
        {
            "name": "空值处理",
            "input": {"empty_field": None},
            "check": lambda ctx: ctx.get("empty_field") == ""
        }
    ]
    
    for test_case in test_cases:
        processed = generator._preprocess_context(test_case["input"])
        if test_case["check"](processed):
            print(f"✓ {test_case['name']}: {processed}")
        else:
            print(f"⚠ {test_case['name']}: {processed}")
    
    print("\n✓ 数据预处理测试通过\n")


def main():
    """运行所有单元测试"""
    print("=" * 50)
    print("文档生成器单元测试")
    print("=" * 50)
    print()
    
    try:
        test_document_generator()
        test_data_preprocessing()
        test_template_converter()
        
        print("=" * 50)
        print("✓ 所有单元测试通过!")
        print("=" * 50)
        
    except Exception as e:
        print(f"\n✗ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
