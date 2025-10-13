"""文档生成服务 - 模块四"""
from docxtpl import DocxTemplate
from typing import Dict, Any, List, Optional
import io
import logging
from datetime import datetime
import re

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Word 文档生成器（使用 Jinja2 模板）"""
    
    def __init__(self, template_bytes: bytes):
        """
        初始化文档生成器
        
        Args:
            template_bytes: 模板文件的字节数据
        """
        try:
            self.template = DocxTemplate(io.BytesIO(template_bytes))
            self.template_size = len(template_bytes)
            logger.info(f"模板初始化成功，大小: {self.template_size} 字节")
        except Exception as e:
            logger.error(f"模板初始化失败: {e}")
            raise ValueError(f"无法加载模板: {str(e)}")
    
    def render(self, context: Dict[str, Any]) -> bytes:
        """
        渲染文档
        
        Args:
            context: 上下文数据字典
            
        Returns:
            生成的文档字节数据
        """
        try:
            # 预处理上下文数据
            processed_context = self._preprocess_context(context)
            
            logger.info(f"开始渲染文档，变量数: {len(processed_context)}")
            
            # 填充数据到模板
            self.template.render(processed_context)
            
            # 保存到内存
            file_stream = io.BytesIO()
            self.template.save(file_stream)
            file_stream.seek(0)
            
            result_bytes = file_stream.getvalue()
            logger.info(f"文档渲染成功，大小: {len(result_bytes)} 字节")
            
            return result_bytes
            
        except Exception as e:
            logger.error(f"文档渲染失败: {e}", exc_info=True)
            raise ValueError(f"文档渲染失败: {str(e)}")
    
    def render_with_tables(
        self, 
        context: Dict[str, Any], 
        table_data: Dict[str, List[Dict]]
    ) -> bytes:
        """
        渲染包含表格数据的文档
        
        Args:
            context: 基础上下文数据
            table_data: 表格数据，格式：{"table_name": [{"col1": "val1", ...}, ...]}
            
        Returns:
            生成的文档字节数据
        """
        # 合并上下文和表格数据
        full_context = {
            **context,
            **table_data
        }
        
        logger.info(f"渲染表格文档，表格数: {len(table_data)}")
        return self.render(full_context)
    
    def _preprocess_context(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        预处理上下文数据
        
        Args:
            context: 原始上下文数据
            
        Returns:
            处理后的上下文数据
        """
        processed = {}
        
        for key, value in context.items():
            if value is None:
                # 空值替换为空字符串
                processed[key] = ""
                continue
            
            # 处理日期格式
            if isinstance(value, str) and self._is_date_string(value):
                processed[key] = self._format_date(value)
                continue
            
            # 处理数字格式（金额）
            if isinstance(value, (int, float)) and any(
                keyword in key.lower() 
                for keyword in ['amount', 'price', 'money', 'fee', '金额', '价格', '费用']
            ):
                processed[key] = self._format_money(value)
                continue
            
            # 处理布尔值
            if isinstance(value, bool):
                processed[key] = "是" if value else "否"
                continue
            
            # 其他类型直接使用
            processed[key] = value
        
        return processed
    
    def _is_date_string(self, value: str) -> bool:
        """判断字符串是否为日期格式"""
        date_patterns = [
            r'^\d{4}-\d{2}-\d{2}$',  # YYYY-MM-DD
            r'^\d{4}/\d{2}/\d{2}$',  # YYYY/MM/DD
            r'^\d{4}\.\d{2}\.\d{2}$',  # YYYY.MM.DD
        ]
        
        for pattern in date_patterns:
            if re.match(pattern, value):
                return True
        
        return False
    
    def _format_date(self, date_string: str) -> str:
        """
        格式化日期
        
        Args:
            date_string: 日期字符串（YYYY-MM-DD 等格式）
            
        Returns:
            中文格式日期（YYYY年MM月DD日）
        """
        try:
            # 尝试多种日期格式
            for fmt in ['%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d']:
                try:
                    date_obj = datetime.strptime(date_string, fmt)
                    return date_obj.strftime('%Y年%m月%d日')
                except ValueError:
                    continue
            
            # 如果都失败，返回原字符串
            return date_string
            
        except Exception as e:
            logger.warning(f"日期格式化失败: {e}")
            return date_string
    
    def _format_money(self, amount: float) -> str:
        """
        格式化金额
        
        Args:
            amount: 金额数字
            
        Returns:
            格式化的金额字符串（带千位分隔符）
        """
        try:
            # 格式化为带千位分隔符的字符串，保留两位小数
            return f"{amount:,.2f}"
        except Exception as e:
            logger.warning(f"金额格式化失败: {e}")
            return str(amount)
    
    @staticmethod
    def convert_placeholders_to_jinja(text: str) -> str:
        """
        将占位符格式从 {{变量}} 转换为 Jinja2 格式 {{ 变量 }}
        
        Args:
            text: 原始文本
            
        Returns:
            转换后的文本
        """
        # 这个函数用于辅助模板转换
        # 实际使用中，模板应该已经是 Jinja2 格式
        return text
    
    @staticmethod
    def validate_template(template_bytes: bytes) -> Dict[str, Any]:
        """
        验证模板文件
        
        Args:
            template_bytes: 模板文件字节数据
            
        Returns:
            验证结果
        """
        try:
            template = DocxTemplate(io.BytesIO(template_bytes))
            
            # 提取模板中的变量
            # docxtpl 会自动解析 Jinja2 变量
            
            return {
                "valid": True,
                "size": len(template_bytes),
                "message": "模板有效"
            }
            
        except Exception as e:
            return {
                "valid": False,
                "error": str(e),
                "message": f"模板验证失败: {str(e)}"
            }


class TemplateConverter:
    """模板转换器（将普通模板转换为 Jinja2 模板）"""
    
    @staticmethod
    def convert_to_jinja_template(
        input_bytes: bytes,
        variable_mapping: Optional[Dict[str, str]] = None
    ) -> bytes:
        """
        将包含 {{变量}} 的普通 Word 文档转换为 Jinja2 模板
        
        Args:
            input_bytes: 输入文档字节数据
            variable_mapping: 变量映射（中文标签 -> 英文变量名）
            
        Returns:
            转换后的模板字节数据
        """
        from docx import Document
        
        try:
            doc = Document(io.BytesIO(input_bytes))
            
            # 如果没有映射，直接返回
            if not variable_mapping:
                logger.info("无变量映射，直接返回原文档")
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                return output.getvalue()
            
            # 替换段落中的占位符
            for para in doc.paragraphs:
                if '{{' in para.text and '}}' in para.text:
                    new_text = para.text
                    for label, var_name in variable_mapping.items():
                        # 将 {{中文标签}} 替换为 {{ 英文变量名 }}
                        new_text = new_text.replace(
                            f"{{{{{label}}}}}",
                            f"{{{{ {var_name} }}}}"
                        )
                    
                    # 清空段落并重新添加文本
                    para.clear()
                    para.add_run(new_text)
            
            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if '{{' in para.text and '}}' in para.text:
                                new_text = para.text
                                for label, var_name in variable_mapping.items():
                                    new_text = new_text.replace(
                                        f"{{{{{label}}}}}",
                                        f"{{{{ {var_name} }}}}"
                                    )
                                para.clear()
                                para.add_run(new_text)
            
            # 保存转换后的文档
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            result_bytes = output.getvalue()
            logger.info(f"模板转换成功，大小: {len(result_bytes)} 字节")
            
            return result_bytes
            
        except Exception as e:
            logger.error(f"模板转换失败: {e}")
            raise ValueError(f"模板转换失败: {str(e)}")


# 全局实例
_document_generator: Optional[DocumentGenerator] = None


def create_document_generator(template_bytes: bytes) -> DocumentGenerator:
    """创建文档生成器实例"""
    return DocumentGenerator(template_bytes)
