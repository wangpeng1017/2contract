"""
模块一：模板解析服务
使用 python-docx 解析 .docx 文件，提取纯文本内容
"""
import logging
from pathlib import Path
from typing import List, Tuple
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class TemplateParser:
    """Word文档模板解析器"""
    
    def __init__(self):
        """初始化解析器"""
        pass
    
    def parse_docx(self, file_path: Path) -> str:
        """
        解析docx文件，提取所有文本内容
        
        Args:
            file_path: docx文件路径
            
        Returns:
            str: 提取的纯文本内容，保持文档阅读顺序
        """
        try:
            logger.info(f"开始解析文档: {file_path}")
            doc = Document(file_path)
            
            # 按照文档顺序提取内容
            text_parts = []
            
            # 遍历文档的所有元素（段落和表格）
            for element in doc.element.body:
                # 处理段落
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)
                
                # 处理表格
                elif element.tag.endswith('tbl'):
                    table = Table(element, doc)
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append(table_text)
            
            # 合并所有文本
            full_text = '\n\n'.join(text_parts)
            
            logger.info(f"文档解析完成，提取了 {len(text_parts)} 个文本块")
            logger.debug(f"提取的文本长度: {len(full_text)} 字符")
            
            return full_text
            
        except Exception as e:
            logger.error(f"解析文档失败: {str(e)}")
            raise ValueError(f"无法解析文档: {str(e)}")
    
    def _extract_table_text(self, table: Table) -> str:
        """
        从表格中提取文本
        
        Args:
            table: docx表格对象
            
        Returns:
            str: 表格文本内容
        """
        table_texts = []
        
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            
            if row_texts:
                # 用制表符分隔单元格
                table_texts.append(' | '.join(row_texts))
        
        return '\n'.join(table_texts)
    
    def extract_paragraphs_and_tables(self, file_path: Path) -> Tuple[List[str], List[List[List[str]]]]:
        """
        分别提取段落和表格数据（用于更精细的处理）
        
        Args:
            file_path: docx文件路径
            
        Returns:
            Tuple[List[str], List[List[List[str]]]]: (段落列表, 表格列表)
        """
        try:
            doc = Document(file_path)
            
            # 提取段落
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            logger.info(f"提取了 {len(paragraphs)} 个段落和 {len(tables)} 个表格")
            
            return paragraphs, tables
            
        except Exception as e:
            logger.error(f"提取段落和表格失败: {str(e)}")
            raise ValueError(f"无法提取文档内容: {str(e)}")
    
    def get_document_metadata(self, file_path: Path) -> dict:
        """
        获取文档元数据
        
        Args:
            file_path: docx文件路径
            
        Returns:
            dict: 文档元数据
        """
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created": core_props.created,
                "modified": core_props.modified,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            
            return metadata
            
        except Exception as e:
            logger.warning(f"获取文档元数据失败: {str(e)}")
            return {}


# 创建全局实例
template_parser = TemplateParser()
