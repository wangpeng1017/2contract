"""文档解析服务 - 模块一"""
from docx import Document
from typing import Dict, List, Any, Optional
import io
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Word 文档解析器"""
    
    def __init__(self, file_bytes: bytes):
        """
        初始化文档解析器
        
        Args:
            file_bytes: 文档的字节数据
        """
        try:
            self.doc = Document(io.BytesIO(file_bytes))
            self.file_size = len(file_bytes)
        except Exception as e:
            logger.error(f"文档初始化失败: {e}")
            raise ValueError(f"无法解析文档: {str(e)}")
    
    def extract_text(self) -> str:
        """
        提取文档全文（保持阅读顺序）
        
        Returns:
            文档的纯文本内容，按照段落和表格在文档中的顺序
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        
        full_text = []
        
        try:
            # 按照文档元素顺序遍历（段落和表格交错）
            for element in self.doc.element.body:
                # 处理段落
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, self.doc)
                    text = paragraph.text.strip()
                    if text:
                        full_text.append(text)
                
                # 处理表格
                elif element.tag.endswith('tbl'):
                    table = Table(element, self.doc)
                    table_text = self._extract_table_text_ordered(table)
                    if table_text:
                        full_text.append(table_text)
            
            result = "\n\n".join(full_text)
            logger.info(f"成功提取文本（阅读顺序），长度: {len(result)} 字符")
            return result
            
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            raise
    
    def get_structure(self) -> Dict[str, Any]:
        """
        获取文档结构信息
        
        Returns:
            文档结构的字典表示
        """
        try:
            # 收集使用的样式
            styles_used = set()
            for para in self.doc.paragraphs:
                if para.style and para.style.name:
                    styles_used.add(para.style.name)
            
            # 分析表格结构
            tables_info = []
            for i, table in enumerate(self.doc.tables):
                tables_info.append({
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0
                })
            
            structure = {
                "paragraphs_count": len(self.doc.paragraphs),
                "tables_count": len(self.doc.tables),
                "sections_count": len(self.doc.sections),
                "styles_used": sorted(list(styles_used)),
                "tables_info": tables_info,
                "file_size_bytes": self.file_size
            }
            
            logger.info(f"文档结构: {structure['paragraphs_count']} 段落, "
                       f"{structure['tables_count']} 表格")
            return structure
            
        except Exception as e:
            logger.error(f"结构分析失败: {e}")
            raise
    
    def validate(self) -> Dict[str, Any]:
        """
        验证文档有效性
        
        Returns:
            验证结果字典
        """
        try:
            text = self.extract_text()
            
            validation_result = {
                "valid": True,
                "text_length": len(text),
                "has_content": len(text) > 0,
                "has_tables": len(self.doc.tables) > 0,
                "has_paragraphs": len(self.doc.paragraphs) > 0,
                "warnings": []
            }
            
            # 检查是否有内容
            if not validation_result["has_content"]:
                validation_result["warnings"].append("文档没有可提取的文本内容")
            
            # 检查文本长度是否合理
            if len(text) < 50:
                validation_result["warnings"].append("文档内容过少，可能不是有效的合同模板")
            
            if len(text) > 100000:
                validation_result["warnings"].append("文档内容过长，处理可能需要较长时间")
            
            logger.info(f"文档验证完成: {validation_result}")
            return validation_result
            
        except Exception as e:
            logger.error(f"文档验证失败: {e}")
            return {
                "valid": False,
                "error": str(e)
            }
    
    def extract_placeholders(self) -> List[str]:
        """
        提取文档中的占位符（格式：{{变量名}}）
        
        Returns:
            占位符列表
        """
        import re
        
        text = self.extract_text()
        # 匹配 {{变量名}} 格式的占位符
        pattern = r'\{\{([^}]+)\}\}'
        placeholders = re.findall(pattern, text)
        
        # 去重并排序
        unique_placeholders = sorted(list(set(placeholders)))
        
        logger.info(f"找到 {len(unique_placeholders)} 个占位符: {unique_placeholders}")
        return unique_placeholders
    
    def get_metadata(self) -> Dict[str, Any]:
        """
        获取文档元数据
        
        Returns:
            文档元数据字典
        """
        try:
            core_properties = self.doc.core_properties
            
            metadata = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "created": core_properties.created.isoformat() if core_properties.created else None,
                "modified": core_properties.modified.isoformat() if core_properties.modified else None,
                "last_modified_by": core_properties.last_modified_by or ""
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"获取元数据失败: {e}")
            return {}
    
    def _extract_table_text_ordered(self, table) -> str:
        """
        从表格中提取文本（保持行列顺序）
        
        Args:
            table: docx表格对象
            
        Returns:
            str: 表格文本内容
        """
        table_texts = []
        
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            
            if row_texts:
                # 用制表符分隔单元格
                table_texts.append(' | '.join(row_texts))
        
        return '\n'.join(table_texts)


class DocumentParserFactory:
    """文档解析器工厂"""
    
    @staticmethod
    def create_parser(file_bytes: bytes, file_extension: str = ".docx") -> DocumentParser:
        """
        创建文档解析器
        
        Args:
            file_bytes: 文件字节数据
            file_extension: 文件扩展名
            
        Returns:
            DocumentParser 实例
        """
        if file_extension.lower() != ".docx":
            raise ValueError(f"不支持的文件格式: {file_extension}")
        
        return DocumentParser(file_bytes)
