"""文档生成路由 - 模块四"""
from fastapi import APIRouter, HTTPException, Response
from pydantic import BaseModel, Field
from typing import Dict, Any, Optional
import logging
from datetime import datetime
import uuid

from services.document_generator import create_document_generator, DocumentGenerator, TemplateConverter
from services.storage_service import get_storage_service
from models.database import get_db
from models.template import Template
from models.generated_document import GeneratedDocument

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/generate", tags=["generate"])


class GenerateRequest(BaseModel):
    """文档生成请求"""
    template_id: str = Field(..., description="模板 ID")
    data: Dict[str, Any] = Field(..., description="填充数据")
    filename: Optional[str] = Field(None, description="输出文件名")
    template_storage_key: Optional[str] = Field(None, description="模板在MinIO中的存储键")


class GenerateResponse(BaseModel):
    """文档生成响应"""
    document_id: str = Field(..., description="生成的文档 ID")
    download_url: str = Field(..., description="下载链接")
    filename: str = Field(..., description="文件名")
    file_size: int = Field(..., description="文件大小（字节）")
    generation_time_ms: float = Field(..., description="生成耗时（毫秒）")


@router.post("/document", response_model=GenerateResponse)
async def generate_document(request: GenerateRequest):
    """
    生成合同文档
    
    - **template_id**: 模板 ID（可以是 "test" 用于测试）
    - **data**: 填充数据字典
    - **filename**: 可选的输出文件名
    
    Returns:
        生成的文档信息和下载链接
    """
    start_time = datetime.now()
    
    try:
        logger.info(f"开始生成文档: template_id={request.template_id}")
        
        # 获取模板内容
        if request.template_storage_key:
            # 从 MinIO 直接下载模板
            storage = get_storage_service()
            template_bytes = await storage.download(request.template_storage_key)
            logger.info(f"从 MinIO 加载模板: {request.template_storage_key}")
        elif request.template_id == "test":
            # 测试模式：使用测试模板
            template_bytes = await _get_test_template()
        else:
            # 从数据库和存储获取模板
            template_bytes = await _get_template_from_storage(request.template_id)
        
        # 创建文档生成器
        generator = create_document_generator(template_bytes)
        
        # 渲染文档
        output_bytes = generator.render(request.data)
        
        # 生成文件名
        output_filename = request.filename or f"contract_{uuid.uuid4().hex[:8]}.docx"
        if not output_filename.endswith('.docx'):
            output_filename += '.docx'
        
        # 上传到 MinIO
        storage = get_storage_service()
        document_id = str(uuid.uuid4())
        output_key = f"generated/{datetime.now().strftime('%Y%m%d')}/{document_id}_{output_filename}"
        
        # 上传文件到 MinIO
        storage_url = await storage.upload(
            object_name=output_key,
            data=output_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={
                "document_id": document_id,
                "original_filename": output_filename,
                "generation_time": datetime.now().isoformat()
            }
        )
        
        # 生成预签名 URL（1小时有效）
        from datetime import timedelta
        download_url = storage.get_presigned_url(
            object_name=output_key,
            expires=timedelta(hours=1)
        )
        
        # 同时保留缓存作为备选方案
        from services.cache_service import get_cache_service
        cache = get_cache_service()
        import base64
        doc_info = {
            "document_id": document_id,
            "filename": output_filename,
            "storage_key": output_key,
            "storage_url": storage_url,
            "file_size": len(output_bytes),
            "file_content": base64.b64encode(output_bytes).decode('utf-8'),  # Base64 编码
            "created_at": datetime.now().isoformat()
        }
        
        await cache.set(f"generated_doc:{document_id}", doc_info, ttl=3600)  # 1小时
        
        # 计算生成时间
        generation_time = (datetime.now() - start_time).total_seconds() * 1000
        
        # TODO: 保存到数据库
        # record = GeneratedDocument(...)
        # db.add(record)
        # db.commit()
        
        logger.info(
            f"文档生成成功: id={document_id}, "
            f"size={len(output_bytes)}, time={generation_time:.2f}ms"
        )
        
        return GenerateResponse(
            document_id=document_id,
            download_url=download_url,  # MinIO 预签名 URL
            filename=output_filename,
            file_size=len(output_bytes),
            generation_time_ms=generation_time
        )
        
    except ValueError as e:
        logger.error(f"文档生成失败: {e}")
        raise HTTPException(400, f"文档生成失败: {str(e)}")
    except Exception as e:
        logger.error(f"文档生成异常: {e}", exc_info=True)
        raise HTTPException(500, f"服务器错误: {str(e)}")


@router.get("/document/{document_id}/download")
async def download_document(document_id: str):
    """
    下载生成的文档（备用接口，主要使用 MinIO 预签名 URL）
    
    - **document_id**: 文档 ID
    """
    try:
        # 从缓存获取文档信息
        from services.cache_service import get_cache_service
        cache = get_cache_service()
        
        doc_info = await cache.get(f"generated_doc:{document_id}")
        
        if not doc_info:
            raise HTTPException(404, "文档不存在或已过期")
        
        # 优先从 MinIO 下载
        if "storage_key" in doc_info:
            try:
                storage = get_storage_service()
                file_bytes = await storage.download(doc_info["storage_key"])
                
                return Response(
                    content=file_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="{doc_info["filename"]}"'
                    }
                )
            except Exception as e:
                logger.warning(f"从 MinIO 下载失败，回退到缓存: {e}")
        
        # 回退到缓存下载
        if "file_content" in doc_info:
            import base64
            file_bytes = base64.b64decode(doc_info["file_content"])
            
            return Response(
                content=file_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="{doc_info["filename"]}"'
                }
            )
        
        raise HTTPException(404, "文档内容不可用")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档下载失败: {e}", exc_info=True)
        raise HTTPException(500, f"文档下载失败: {str(e)}")


@router.post("/template/convert")
async def convert_template(
    file: bytes,
    variable_mapping: Optional[Dict[str, str]] = None
):
    """
    转换模板格式
    
    将包含 {{变量}} 的普通文档转换为 Jinja2 模板
    
    - **file**: Word 文档文件
    - **variable_mapping**: 变量映射（可选）
    """
    try:
        converter = TemplateConverter()
        converted_bytes = converter.convert_to_jinja_template(
            input_bytes=file,
            variable_mapping=variable_mapping
        )
        
        return Response(
            content=converted_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": 'attachment; filename="converted_template.docx"'
            }
        )
        
    except Exception as e:
        logger.error(f"模板转换失败: {e}")
        raise HTTPException(500, f"模板转换失败: {str(e)}")


@router.post("/template/validate")
async def validate_template(file: bytes):
    """
    验证模板文件
    
    - **file**: Word 文档文件
    """
    try:
        result = DocumentGenerator.validate_template(file)
        return result
    except Exception as e:
        logger.error(f"模板验证失败: {e}")
        raise HTTPException(500, f"模板验证失败: {str(e)}")


# 辅助函数

async def _get_test_template() -> bytes:
    """获取测试模板"""
    # 创建一个简单的测试模板
    from docx import Document
    
    doc = Document()
    doc.add_heading('采购合同', 0)
    doc.add_paragraph()
    
    # 添加合同内容
    doc.add_paragraph(f'合同编号：{{{{ contract_number }}}}')
    doc.add_paragraph(f'签订日期：{{{{ signing_date }}}}')
    doc.add_paragraph()
    doc.add_paragraph(f'甲方：{{{{ party_a }}}}')
    doc.add_paragraph(f'联系人：{{{{ party_a_contact }}}}')
    doc.add_paragraph(f'电话：{{{{ party_a_phone }}}}')
    doc.add_paragraph()
    doc.add_paragraph(f'乙方：{{{{ party_b }}}}')
    doc.add_paragraph(f'联系人：{{{{ party_b_contact }}}}')
    doc.add_paragraph(f'电话：{{{{ party_b_phone }}}}')
    doc.add_paragraph()
    doc.add_paragraph(f'合同金额：{{{{ contract_amount }}}} 元')
    doc.add_paragraph(f'付款方式：{{{{ payment_method }}}}')
    doc.add_paragraph(f'交付时间：{{{{ delivery_time }}}}')
    doc.add_paragraph(f'交付地址：{{{{ delivery_address }}}}')
    
    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()


async def _get_template_from_storage(template_id: str) -> bytes:
    """从存储获取模板"""
    # TODO: 从数据库查询模板记录
    # TODO: 从 MinIO 下载模板文件
    
    # 简化实现：使用测试模板
    logger.warning(f"模板 {template_id} 不存在，使用测试模板")
    return await _get_test_template()


import io
